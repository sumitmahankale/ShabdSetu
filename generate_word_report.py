import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_from_markdown(doc, table_md):
    """Parses a Markdown table and adds it to the document."""
    lines = [line.strip() for line in table_md.strip().split('\n')]
    
    # Extract headers
    headers = [header.strip() for header in lines[0].strip('|').split('|')]
    num_cols = len(headers)
    
    # Data rows
    rows_data = []
    for line in lines[2:]: # Skip separator line
        rows_data.append([cell.strip() for cell in line.strip('|').split('|')])

    # Create table
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make header bold
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    # Add data rows
    for row_data in rows_data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

def markdown_to_docx(md_content, docx_filename):
    """Converts Markdown content to a DOCX file."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Split content by lines
    lines = md_content.strip().split('\n')
    
    in_table = False
    table_md = ""

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_md = ""
            table_md += line + '\n'
            continue
        elif in_table:
            add_table_from_markdown(doc, table_md)
            in_table = False
            table_md = ""

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Lists (simple unordered)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # Bold/Italic text in paragraphs
        else:
            p = doc.add_paragraph()
            # Split by bold/italic markers
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                else:
                    p.add_run(part)

    doc.save(docx_filename)
    print(f"Successfully created {docx_filename}")

if __name__ == '__main__':
    try:
        with open('comparison_report.md', 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        output_filename = 'ShabdSetu_Comparison_Report.docx'
        markdown_to_docx(markdown_content, output_filename)
        
    except FileNotFoundError:
        print("Error: comparison_report.md not found.")
    except Exception as e:
        print(f"An error occurred: {e}")
